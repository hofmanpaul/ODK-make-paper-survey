#This script takes all ODK formatted excel files in the current directory and turns them into paper surveys
#In no way is this supposed to lead to a perfect file. Extensive tweaking will be necessary, as 
# the structure of ODK allows for completely different surveywriting.
#Python 3.6 and up

from docx import Document
import pandas
import re
import os

def replace_dollarrefs(text, type):
    #Keep going until all ODK references are replaced with human-readeable references
    while '$' in text:
    #find the first relevant variable
        match=re.search('\{', text)
        first=match.start()+1
        match=re.search('}', text)
        last=match.end()-1
        variable=text[first:last]
        #Remove the curly braces, and add the question number
        Found=False
        i=0
        while not Found:
            if numbered_varlist[i][1]==variable:
                text=text.replace('${' + variable + '}', '[Q' + str(numbered_varlist[i][0]) + ' ' + numbered_varlist[i][1]  + ']')
                Found=True
            i+=1
        #Additional replaces for selected syntax in relevants
        while 'selected' in text and type=='relevant':
            match=re.search('selected\(.+?\)', text)
            relevant_orig=text[match.start():match.end()]
            relevant_new=relevant_orig
            relevant_new=relevant_new.replace('selected(', '')
            relevant_new=relevant_new.replace(',', '=')
            relevant_new=relevant_new.replace("'", '')
            relevant_new=relevant_new.replace(")", '')
            text=text.replace(relevant_orig, relevant_new)
    return text

# Initialise Files
listfiles=os.listdir(os.getcwd())
excelfiles=[]
for filename in listfiles:
    if filename.endswith('.xls') or filename.endswith('.xlsx') and ('~$') not in filename:
        excelfiles.append(filename)
print("This file will work on:") 
for excelfile in excelfiles:
    print(excelfile)
for excelfile in excelfiles:
    survey= pandas.read_excel(excelfile, sheet_name='survey')
    choices= pandas.read_excel(excelfile, sheet_name='choices')
    settings= pandas.read_excel(excelfile, sheet_name='settings')
    outdoc=Document()
    try:
        language = "::" + settings.at[0,'default language'].strip().lower()
        print(f"Language found: {language} ")
    except Exception:
        language = ""
        print("No language found")
    label_col = 'label' + language
    hint_col = 'hint' + language
    if 'relevant' in survey.columns:
        relevant_col = 'relevant'
    elif 'relevance' in survey.columns:
        relevant_col = 'relevance'
    else:
        print('Relevance column not found')
        
    choicesdict={}
    choicelist=[]
    listname=choices.at[0, 'list_name'].strip().lower()
    for index, row in choices.iterrows():
        if row['list_name']=="":
            continue
        if listname==row['list_name'].strip().lower(): 
            choicelist.append(str(row['name']) + '. ' + row['label'])
        else: #write the choicelist to the dictionary, then continue with the next choices
            choicesdict[listname]=choicelist
            choicelist=[]
            listname=row['list_name'].strip().lower()
            choicelist.append(str(row['name']) + '. ' + row['label'])
    choicesdict[listname]=choicelist #for the last choice list


    #Write Initial stuff
    outdoc.add_heading(settings.at[0, 'form_title'])
    outdoc.add_paragraph("This is an automatically generated paper survey based on an ODK excel " 
                         "file. Each question consists of four parts: 1. the variable name (bold). " 
                         "This also includes a question number (which is not used) and matches the " 
                         "name of the corresponding variable in the dataset. 2. A label. 3. a hint " 
                         "(italic). 4. The answer option(s). Before the hint there can be some " 
                         "conditions that dictate when this question is asked. The answer options " 
                         "are: ellipses (...) for text answers, white space between hyphens (-    -) "
                         "for numbers, o if interviewees have to choose one and u if interviewees " 
                         "have to select multiple.")


    skiplist = ['begin group', 'end group', 'start', 'end', 'deviceid','begin repeat', 'end repeat'
                , 'note']
    # Make list of lists of the question properties and write
    numbered_varlist=[]
    questionnumber=1
    for index, row in survey.iterrows():
        if row['type'] in skiplist:
            continue
        
        if ' ' in row['type'].strip():
            typelist=row['type'].strip().split(' ')
            qtype=typelist[0]
            choices=typelist[1]
        else:
            qtype=row['type']
            choices=''
        #Write!
        #Name
        namewrite=outdoc.add_paragraph()
        namewrite.add_run(str(questionnumber)+ ' ' + row['name']).bold=True
        numbered_varlist.append([questionnumber, row['name']])
        
        #Relevant
        if pandas.notnull(row[relevant_col]):
            relevant = replace_dollarrefs(row[relevant_col], 'relevant')
            relevantwrite='Only ask if ' + relevant
            relevantwrite=outdoc.add_paragraph(relevantwrite)      
            
        #Label
        if pandas.notnull(row[label_col]):
            label = replace_dollarrefs(row[label_col], 'label')
            labelwrite=outdoc.add_paragraph()
            labelwrite.add_run(label)
    
        #Hint
        if pandas.notnull(row[hint_col]):
            hint = replace_dollarrefs(row[hint_col], 'hint')
            hintwrite=outdoc.add_paragraph()
            hintwrite.add_run(hint).italic=True
    
        #Choices
        if qtype=="text" or qtype=="string":
            columnwrite=outdoc.add_paragraph("...................................................."
                                             "...................................................."
                                             ".........................................")
        if qtype=="integer" or qtype=="decimal":
            columnwrite=outdoc.add_paragraph('-             -')
        if qtype=='select_one' or qtype=='select_multiple':
            if qtype=='select_one':
                bullet='o'
            elif qtype=='select_multiple':
                bullet="u"
            else:
                for option in choicesdict[choices]:
                    optionwrite=outdoc.add_paragraph(bullet + ' ' + option)

        questionnumber+=1

    outdoc.save(excelfile+'.docx')
print('Finished')
